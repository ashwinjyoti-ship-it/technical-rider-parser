from flask import Flask, render_template, request, jsonify, send_file
import os
import re
from datetime import datetime
from werkzeug.utils import secure_filename
import pdfplumber
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = '/tmp/uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx', 'doc', 'txt'}

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Equipment brand databases
MICROPHONE_BRANDS = {
    'shure': ['shure', 'sm57', 'sm58', 'beta', 'ksm', 'axient', 'ulxd', 'psm'],
    'sennheiser': ['sennheiser', 'e901', 'e902', 'e904', 'e914', 'e614', 'd6000', 'mm445'],
    'akg': ['akg', 'c411', 'c414', 'c451', 'c12'],
    'dpa': ['dpa', '4099', '4011', '4021'],
    'audix': ['audix', 'd6', 'i5', 'om5'],
    'neumann': ['neumann', 'u87', 'u89', 'km184', 'km81'],
    'rode': ['rode', 'nt1', 'nt5', 'ntg'],
    'earthworks': ['earthworks', 'sr25', 'dm20'],
    'beyerdynamic': ['beyerdynamic', 'tg', 'opus'],
    'audio-technica': ['audio-technica', 'at', 'at4050'],
    'electrovoice': ['electrovoice', 're20', 'nd68'],
}

DI_BRANDS = ['radial', 'bss', 'klarkteknik', 'klark', 'countryman', 'avalon', 'whirlwind']

MONITOR_BRANDS = ["l'acoustic", 'lacoustic', 'meyer sound', 'meyer', 'd&b', 'dnb', 'db', 
                  'outline', 'das', 'jbl', 'yamaha', 'qsc', 'electro-voice', 'ev']

IEM_BRANDS = ['sennheiser', 'shure', 'psm', 'psm1000', 'psm900', 'ew300', '2050', 'g4', 
              'ulxd', 'axient']

BACKLINE_KEYWORDS = {
    'drums': ['drum', 'kit', 'drumkit', 'pearl', 'tama', 'dw', 'yamaha drums', 
              'kick', 'snare', 'tom', 'cymbal', 'hi-hat', 'hihat', 'throne', 
              'pedal', 'zildjian', 'sabian', 'meinl', 'paiste'],
    'keyboards': ['keyboard', 'piano', 'synthesizer', 'synth', 'yamaha', 'nord', 
                  'korg', 'roland', 'moog', 'montage', 'stage', 'stand', 'sustain pedal'],
    'guitar': ['guitar', 'bass', 'amplifier', 'amp', 'cabinet', 'fender', 'marshall', 
               'mesa', 'vox', 'orange', 'ampeg', 'combo', 'head'],
    'other': ['riser', 'stand', 'chair', 'stool', 'cable', 'power']
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def extract_text_from_pdf(file_path):
    """Extract text from PDF using pdfplumber"""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        print(f"Error extracting PDF: {e}")
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX"""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
    return text

def extract_text_from_txt(file_path):
    """Extract text from TXT"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        print(f"Error extracting TXT: {e}")
        return ""

def parse_microphones(text):
    """Parse microphones - simple brand and model extraction"""
    text_lower = text.lower()
    mic_list = {}
    
    # Specific model patterns - only what's actually in the document
    specific_patterns = [
        # Samson
        (r'(\d+)\s*(?:x|nos\.?).*?samson\s+co2', 'Samson CO2'),
        
        # Shure models
        (r'(\d+)\s*(?:x|nos\.?).*?(?:shure\s+)?sm58', 'Shure SM58'),
        (r'(\d+)\s*(?:x|nos\.?).*?(?:shure\s+)?sm57', 'Shure SM57'),
        (r'(\d+)\s*(?:x|nos\.?).*?(?:shure\s+)?beta\s*58', 'Shure Beta58'),
        (r'(\d+)\s*(?:x|nos\.?).*?(?:shure\s+)?beta\s*81', 'Shure Beta81'),
        (r'(\d+)\s*(?:x|nos\.?).*?(?:shure\s+)?beta\s*91', 'Shure Beta91'),
        (r'(\d+)\s*(?:x|nos\.?).*?(?:shure\s+)?ksm', 'Shure KSM'),
        (r'(\d+)\s*(?:x|nos\.?).*?(?:shure\s+)?axient', 'Shure Axient'),
        
        # Sennheiser models
        (r'(\d+)\s*(?:x|nos\.?).*?sennheiser\s+e901', 'Sennheiser e901'),
        (r'(\d+)\s*(?:x|nos\.?).*?sennheiser\s+e902', 'Sennheiser e902'),
        (r'(\d+)\s*(?:x|nos\.?).*?sennheiser\s+e904', 'Sennheiser e904'),
        (r'(\d+)\s*(?:x|nos\.?).*?sennheiser\s+e908', 'Sennheiser e908'),
        (r'(\d+)\s*(?:x|nos\.?).*?sennheiser\s+e914', 'Sennheiser e914'),
        (r'(\d+)\s*(?:x|nos\.?).*?sennheiser\s+e614', 'Sennheiser e614'),
        (r'(\d+)\s*(?:x|nos\.?).*?sennheiser\s+e835', 'Sennheiser e835'),
        (r'(\d+)\s*(?:x|nos\.?).*?sennheiser\s+e945', 'Sennheiser e945'),
        (r'(\d+)\s*(?:x|nos\.?).*?sennheiser\s+e965', 'Sennheiser e965'),
        (r'(\d+)\s*(?:x|nos\.?).*?sennheiser\s+6000', 'Sennheiser 6000'),
        
        # DPA
        (r'(\d+)\s*(?:x|nos\.?).*?dpa\s+4099', 'DPA 4099'),
        
        # AKG
        (r'(\d+)\s*(?:x|nos\.?).*?akg\s+c411', 'AKG C411'),
        (r'(\d+)\s*(?:x|nos\.?).*?akg\s+c414', 'AKG C414'),
        
        # Audix
        (r'(\d+)\s*(?:x|nos\.?).*?audix\s+d6', 'Audix D6'),
        (r'(\d+)\s*(?:x|nos\.?).*?audix\s+adx51', 'Audix ADX51'),
        
        # Neumann
        (r'(\d+)\s*(?:x|nos\.?).*?(?:neumann\s+)?km184', 'Neumann KM184'),
        (r'(\d+)\s*(?:x|nos\.?).*?(?:neumann\s+)?km81', 'Neumann KM81'),
        
        # Audio Technica
        (r'(\d+)\s*(?:x|nos\.?).*?(?:audio-technica|at)\s+pro35', 'Audio-Technica Pro35'),
        
        # Electro-Voice
        (r'(\d+)\s*(?:x|nos\.?).*?(?:electro-voice|ev)\s+pl80', 'Electro-Voice PL80'),
        (r'(\d+)\s*(?:x|nos\.?).*?(?:electro-voice|ev)\s+pl24', 'Electro-Voice PL24'),
    ]
    
    for pattern, model_name in specific_patterns:
        matches = re.finditer(pattern, text_lower)
        for match in matches:
            try:
                qty = int(match.group(1))
                if qty <= 50:  # Sanity check
                    if model_name not in mic_list:
                        mic_list[model_name] = 0
                    mic_list[model_name] += qty
            except:
                pass
    
    # Count individual mic entries from tables (like "Choir Mic 1", "Choir Mic 2", etc.)
    mic_entries = re.findall(r'(?:choir|vocal)\s+mic\s+\d+', text_lower)
    if len(mic_entries) > 0:
        # For each entry, look for what brand is specified nearby
        for entry in mic_entries:
            entry_pos = text_lower.find(entry)
            context = text_lower[entry_pos:min(entry_pos+300, len(text_lower))]
            
            # Check which brand/model is mentioned
            if 'samson' in context or 'co2' in context:
                mic_list['Samson CO2'] = mic_list.get('Samson CO2', 0) + 1
            elif 'beta81' in context or 'beta 81' in context:
                mic_list['Shure Beta81'] = mic_list.get('Shure Beta81', 0) + 1
            elif 'sm58' in context:
                mic_list['Shure SM58'] = mic_list.get('Shure SM58', 0) + 1
            elif 'beta58' in context:
                mic_list['Shure Beta58'] = mic_list.get('Shure Beta58', 0) + 1
            elif 'e835' in context:
                mic_list['Sennheiser e835'] = mic_list.get('Sennheiser e835', 0) + 1
            elif 'e945' in context:
                mic_list['Sennheiser e945'] = mic_list.get('Sennheiser e945', 0) + 1
            elif 'audix' in context or 'adx' in context:
                mic_list['Audix ADX51'] = mic_list.get('Audix ADX51', 0) + 1
            elif 'km184' in context:
                mic_list['Neumann KM184'] = mic_list.get('Neumann KM184', 0) + 1
    
    return mic_list

def parse_di_boxes(text):
    """Parse DI boxes - simple count"""
    text_lower = text.lower()
    di_count = 0
    di_brand = None
    
    # Look for DI box counts
    patterns = [
        r'(\d+)\s*(?:x|nos\.?)\s+(?:active\s+)?di\s+(?:box|channel)',
        r'di\s+(?:box|channel)[es]*[:\s]+(\d+)',
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, text_lower)
        for match in matches:
            try:
                qty = int(match.group(1))
                if qty <= 20:
                    di_count = max(di_count, qty)
                    
                    # Only capture brand if explicitly mentioned right next to count
                    context = text_lower[max(0, match.start()-50):match.end()+100]
                    if 'radial' in context:
                        di_brand = 'Radial'
                    elif 'bss' in context:
                        di_brand = 'BSS'
                    elif 'klarkteknik' in context or 'klark' in context:
                        di_brand = 'Klarkteknik'
            except:
                pass
    
    return {'total': di_count, 'brand': di_brand}

def parse_floor_monitors(text):
    """Parse floor monitors - simple count"""
    text_lower = text.lower()
    monitor_count = 0
    monitor_brand = None
    
    patterns = [
        r'(\d+)\s*(?:x|nos\.?)\s+(?:stage\s+)?(?:wedge|monitor)',
        r'total\s+of\s+(?:six|five|four|three|two|eight|ten)\s+\((\d+)\)\s+monitor',
        r'(\d+)\s+as\s+stage\s+wedges',
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, text_lower)
        for match in matches:
            try:
                qty = int(match.group(1))
                if 1 <= qty <= 20:
                    monitor_count = max(monitor_count, qty)
                    
                    # Only capture brand if clearly stated
                    context = text_lower[max(0, match.start()-100):match.end()+200]
                    if "d&b" in context or "d & b" in context:
                        monitor_brand = "d&b"
                    elif "l'acoustic" in context or "lacoustic" in context:
                        monitor_brand = "L'Acoustic"
                    elif "meyer sound" in context or "meyer" in context:
                        monitor_brand = "Meyer Sound"
            except:
                pass
    
    return {'total': monitor_count, 'brand': monitor_brand}

def parse_iems(text):
    """Parse In-Ear Monitors - simple count"""
    text_lower = text.lower()
    iem_count = 0
    iem_brand = None
    
    patterns = [
        r'(\d{1,2})\s*(?:x|nos\.?)\s+(?:in[\s-]?ear[\s-]?monitor|iem)',
        r'(?:ten|eight|nine|seven|six|five)\s+\((\d{1,2})\).*?(?:sennheiser|shure|psm).*?(?:wireless\s+)?iem',
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, text_lower)
        for match in matches:
            try:
                qty = int(match.group(1))
                if 1 <= qty <= 20:
                    iem_count = max(iem_count, qty)
                    
                    # Check for brand
                    context = text_lower[max(0, match.start()-100):match.end()+200]
                    if 'sennheiser' in context:
                        iem_brand = 'Sennheiser'
                    elif 'shure' in context or 'psm' in context:
                        iem_brand = 'Shure'
            except:
                pass
    
    return {'total': iem_count, 'brand': iem_brand}

def parse_backline(text):
    """Parse backline equipment - only what's explicitly mentioned"""
    text_lower = text.lower()
    backline_items = []
    seen_items = set()  # Avoid duplicates
    
    # Look for specific instruments mentioned
    instrument_patterns = [
        (r'(\d+)\s*(?:x|nos\.?)?\s*vibraphone', 'Vibraphone'),
        (r'vibraphone', 'Vibraphone'),  # Catch without quantity
        (r'(\d+)\s*(?:x|nos\.?)?\s*sarangi', 'Sarangi'),
        (r'sarangi', 'Sarangi'),
        (r'(\d+)\s*(?:x|nos\.?)?\s*(?:yamaha\s+)?montage', 'Keyboard (Yamaha Montage)'),
        (r'(\d+)\s*(?:x|nos\.?)?\s*(?:yamaha\s+)?keyboard', 'Keyboard (Yamaha)'),
        (r'(\d+)\s*(?:x|nos\.?)?\s*(?:nord|korg|roland)\s*(?:keyboard|piano|stage)', 'Keyboard'),
        (r'nord\s+(?:stage|piano|electro)', 'Keyboard (Nord)'),
        (r'(\d+)\s*(?:x|nos\.?)?\s*drum\s*(?:kit|set)', 'Drum Kit'),
        (r'(\d+)\s*(?:pc|piece)\s*drum\s*kit', 'Drum Kit'),
    ]
    
    for pattern, instrument_name in instrument_patterns:
        matches = re.finditer(pattern, text_lower)
        for match in matches:
            # Avoid duplicates
            if instrument_name in seen_items:
                continue
                
            try:
                # Try to extract quantity
                if match.groups() and match.group(1):
                    qty_text = match.group(1)
                    qty = int(qty_text) if qty_text.isdigit() else 1
                else:
                    qty = 1
                    
                if qty <= 10:
                    backline_items.append({
                        'item': instrument_name,
                        'quantity': qty
                    })
                    seen_items.add(instrument_name)
            except:
                # If no quantity found, assume 1
                if instrument_name not in seen_items:
                    backline_items.append({
                        'item': instrument_name,
                        'quantity': 1
                    })
                    seen_items.add(instrument_name)
    
    return backline_items

def create_word_report(data, original_filename):
    """Create Word document report in the specified format"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('TECHNICAL EQUIPMENT LIST')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Equipment Breakdown - Parsed from Technical Rider')
    subtitle_run.font.size = Pt(12)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f'Source Document: {original_filename} | Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    meta_run.font.size = Pt(9)
    meta_run.italic = True
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Create table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    headers = ['CATEGORY', 'EQUIPMENT', 'QTY']
    
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        # Bold header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
        # Shade header
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D9E2F3')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Add microphone data
    if data.get('microphones'):
        for model, qty in data['microphones'].items():
            # Handle both old format (dict with 'count') and new format (direct number)
            count = qty if isinstance(qty, int) else qty.get('count', 0)
            row_cells = table.add_row().cells
            row_cells[0].text = 'MICROPHONES'
            row_cells[1].text = model
            row_cells[2].text = str(count)
    
    # Add DI boxes
    if data.get('di_boxes', {}).get('total', 0) > 0:
        row_cells = table.add_row().cells
        row_cells[0].text = 'DI BOXES'
        brand = data['di_boxes'].get('brand', '')
        if brand:
            row_cells[1].text = f'DI Boxes ({brand})'
        else:
            row_cells[1].text = 'DI Boxes'
        row_cells[2].text = str(data['di_boxes']['total'])
    
    # Add floor monitors
    if data.get('floor_monitors', {}).get('total', 0) > 0:
        row_cells = table.add_row().cells
        row_cells[0].text = 'FLOOR MONITORS'
        brand = data['floor_monitors'].get('brand', '')
        if brand:
            row_cells[1].text = f'Floor Monitors ({brand})'
        else:
            row_cells[1].text = 'Floor Monitors'
        row_cells[2].text = str(data['floor_monitors']['total'])
    
    # Add IEMs
    if data.get('iems', {}).get('total', 0) > 0:
        row_cells = table.add_row().cells
        row_cells[0].text = 'IN-EAR MONITORS'
        brand = data['iems'].get('brand', '')
        if brand:
            row_cells[1].text = f'In-Ear Monitors ({brand})'
        else:
            row_cells[1].text = 'In-Ear Monitors'
        row_cells[2].text = str(data['iems']['total'])
    
    # Add backline items
    if data.get('backline') and isinstance(data['backline'], list):
        for item in data['backline']:
            row_cells = table.add_row().cells
            row_cells[0].text = 'BACKLINE'
            row_cells[1].text = item.get('item', 'Unknown')
            row_cells[2].text = str(item.get('quantity', 1))
    
    # Save to BytesIO
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type. Please upload PDF, DOCX, or TXT'}), 400
    
    try:
        # Save file
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Extract text based on file type
        ext = filename.rsplit('.', 1)[1].lower()
        if ext == 'pdf':
            text = extract_text_from_pdf(filepath)
        elif ext in ['docx', 'doc']:
            text = extract_text_from_docx(filepath)
        else:
            text = extract_text_from_txt(filepath)
        
        # Parse equipment
        microphones = parse_microphones(text)
        di_boxes = parse_di_boxes(text)
        floor_monitors = parse_floor_monitors(text)
        iems = parse_iems(text)
        backline = parse_backline(text)
        
        # Prepare response data
        data = {
            'filename': filename,
            'microphones': microphones,
            'di_boxes': di_boxes,
            'floor_monitors': floor_monitors,
            'iems': iems,
            'backline': backline
        }
        
        # Clean up
        os.remove(filepath)
        
        return jsonify(data)
    
    except Exception as e:
        return jsonify({'error': f'Error processing file: {str(e)}'}), 500

@app.route('/download', methods=['POST'])
def download_report():
    try:
        data = request.json
        filename = data.get('filename', 'document')
        
        # Create Word document
        doc_io = create_word_report(data, filename)
        
        # Generate output filename
        output_filename = f"Equipment_List_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=output_filename
        )
    
    except Exception as e:
        return jsonify({'error': f'Error creating report: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
